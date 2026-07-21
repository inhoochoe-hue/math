from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def save_json(data: Any, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _set_normal_font(doc: Document, size: int = 10) -> None:
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(size)


def save_exam_docx(records: list[dict[str, Any]], path: Path, title: str, academy: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_normal_font(doc, 10)
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if academy:
        p = doc.add_paragraph(academy)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("이름: ____________________    점수: __________")

    number = 1
    for record in records:
        if not record.get("quality", {}).get("passed"):
            continue
        item = record["variant"]
        level = item.get("difficulty", "")
        unit = record.get("analysis", {}).get("unit", "")
        doc.add_heading(f"{number}. [{level}] {item.get('title', unit)}", level=2)
        doc.add_paragraph(item.get("question", ""))
        for choice in item.get("choices", []):
            doc.add_paragraph(str(choice))
        doc.add_paragraph("\n")
        number += 1
    doc.save(path)


def save_answer_docx(records: list[dict[str, Any]], path: Path, title: str, academy: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_normal_font(doc, 10)
    heading = doc.add_heading(f"{title} - 정답 및 해설", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if academy:
        p = doc.add_paragraph(academy)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    number = 1
    for record in records:
        if not record.get("quality", {}).get("passed"):
            continue
        item = record["variant"]
        analysis = record.get("analysis", {})
        doc.add_heading(f"{number}. 정답: {item.get('answer', '')}", level=2)
        doc.add_paragraph(f"단원: {analysis.get('unit', '')} / 유형: {analysis.get('concept', '')}")
        doc.add_paragraph(f"난이도: {item.get('difficulty', '')} — {item.get('difficulty_reason', '')}")
        doc.add_paragraph(item.get("solution", ""))
        quality = record.get("quality", {})
        doc.add_paragraph(f"자동 품질점수: {quality.get('score', 0)}점")
        number += 1
    doc.save(path)


def save_docx(results: list[dict[str, Any]], path: Path) -> None:
    """구형 CLI 결과 호환용 단일 문서 저장."""
    normalized = []
    for result in results:
        for item in result.get("variants", []):
            normalized.append({
                "analysis": result.get("analysis", {}),
                "variant": item,
                "quality": item.get("validation", {"passed": True, "score": 100}),
            })
    save_exam_docx(normalized, path, "중등 수학 응용문제")
